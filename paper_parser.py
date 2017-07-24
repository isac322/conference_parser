#!/usr/bin/env python3
# coding: UTF-8

import argparse
import json
import os
import sys
from typing import Union
from urllib.parse import ParseResult, urlparse

import docx
from docx.document import Document
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from parser.abstract_parser import BaseParser
from parser.json_encoder import PaperEncoder
from parser.usenix import UsenixParser

SUPPORTED_TYPE = {'json', 'docx'}


def classify(url: str) -> BaseParser:
    result: ParseResult = urlparse(url)

    if result.hostname == 'www.usenix.org':
        return UsenixParser(url)
    else:
        raise NotImplementedError('{} is unrecognizable.'.format(url))


def save_json(parser: BaseParser, target: Union[str, bytes, os.PathLike, None], on_stdout: bool) -> None:
    page_title = parser.title
    result = parser.parse()

    if on_stdout:
        print(json.dumps(result, indent=True, ensure_ascii=False, cls=PaperEncoder))

    else:
        with open(os.path.join(target, page_title + '.json'), 'w') as fp:
            json.dump(result, fp, indent=True, ensure_ascii=False, cls=PaperEncoder)


def save_docx(parser: BaseParser, target: str) -> None:
    document: Document = docx.Document()
    document.add_heading(parser.title, 0)

    result = parser.parse()

    for section in result:
        document.add_heading(section.name)

        for paper in section.papers:
            document.add_heading(paper.title, level=2)

            for p in paper.abstractions:
                p_para: Paragraph = document.add_paragraph(p)
                p_para.paragraph_format.left_indent = Cm(1)
                p_para.paragraph_format.first_line_indent = Cm(0.25)

                p_para.runs[0].font.size = Pt(9)

        document.add_page_break()

    document.save(os.path.join(target, '{}.docx'.format(parser.title)))


def main():
    arg_parser = argparse.ArgumentParser(description='Parse given URL about paper list')
    arg_parser.add_argument('url', metavar='URL', type=str, help="workshop or conference program's URL")
    arg_parser.add_argument('-t', '--type', metavar='TYPE', type=str, choices=SUPPORTED_TYPE,
                            help='output format. {}'.format(' or '.join(SUPPORTED_TYPE)), default='json')

    group = arg_parser.add_mutually_exclusive_group()
    group.add_argument('-p', '--target-path', metavar='TARGET_PATH', dest='target', type=str,
                       help="only accepts path to save result file")
    group.add_argument('--stdout', action='store_true',
                       help="if TYPE is 'json', output to stdout instead of saving the result to a file. "
                            "if TYPE is 'docx', ignored")

    args = arg_parser.parse_args()

    if args.target is None and not args.stdout:
        if args.type != 'json':
            arg_parser.print_usage()
            print('{}: error: argument -t/--type {}: need -p/--target-path argument'.format(__file__, args.type),
                  file=sys.stderr)
            exit(2)

        args.stdout = True

    parser = classify(args.url)

    if args.type == 'json':
        save_json(parser, args.target, args.stdout)
    elif args.type == 'docx':
        save_docx(parser, args.target)
    else:
        raise NotImplementedError('{} is not in {}'.format(args.type, SUPPORTED_TYPE))


if __name__ == '__main__':
    main()
